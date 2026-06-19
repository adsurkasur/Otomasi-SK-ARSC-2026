import pandas as pd
from docx import Document
from docx.shared import Pt
import sys
import os
import argparse
import datetime
import difflib

# BASE_DIR is set to the parent directory (project root) because the script is in scripts/
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

BIRO_BIDANG_MAP = {
    'ADKEU': 'Administrasi dan Keuangan (ADKEU)',
    'INFOKOM': 'Informasi dan Komunikasi (INFOKOM)',
    'PENKOM': 'Penulisan dan Kompetisi (PENKOM)',
    'PSDM': 'Pengembangan Sumber Daya Mahasiswa (PSDM)',
    'RISTEK': 'Riset dan Teknologi (RISTEK)'
}

DEPT_MAP = {
    'IPABIO': 'Ilmu Pangan dan Bioteknologi',
    'TIP': 'Teknologi Industri Pertanian',
    'TB': 'Teknik Biosistem',
    'TEP': 'Teknik Biosistem',
    'ILMU PANGAN DAN BIOTEKNOLOGI': 'Ilmu Pangan dan Bioteknologi',
    'TEKNOLOGI INDUSTRI PERTANIAN': 'Teknologi Industri Pertanian',
    'TEKNIK BIOSISTEM': 'Teknik Biosistem'
}

def custom_capitalize(word):
    for i, char in enumerate(word):
        if char.isalpha():
            return word[:i] + char.upper() + word[i+1:].lower()
    return word.lower()

def proper_title_case(text):
    if pd.isna(text) or not text:
        return ""
    words = str(text).split()
    conjunctions = ['dan', 'atau', 'di', 'ke', 'dari', 'yang', 'untuk', 'dengan']
    new_words = []
    for i, w in enumerate(words):
        lw = w.lower()
        if i > 0 and lw in conjunctions:
            new_words.append(lw)
        else:
            if "-" in w:
                new_words.append("-".join([custom_capitalize(part) for part in w.split("-")]))
            else:
                new_words.append(custom_capitalize(w))
    return " ".join(new_words)

def normalize_dept(dept):
    if pd.isna(dept) or not dept:
        return ""
    d = str(dept).strip().upper()
    if d in DEPT_MAP:
        return DEPT_MAP[d]
    return proper_title_case(str(dept))

def set_run_font(run, font_name='Times New Roman', font_size=12):
    run.font.name = font_name
    run.font.size = Pt(font_size)

def set_paragraph_text(p, new_text):
    p.text = ""
    run = p.add_run(new_text)
    set_run_font(run)

def get_roman_month(month):
    romans = ["", "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI", "XII"]
    return romans[month]

def create_docx_for_person(person, doc_number, template_doc_path=None):
    if template_doc_path is None:
        template_doc_path = os.path.join(BASE_DIR, 'templates', 'Template_SK.docx')
    # Normalize fields
    nama = proper_title_case(person['Nama Lengkap'])
    nim = str(person['NIM'])
    departemen = normalize_dept(person['Departemen'])
    program_studi = proper_title_case(person['Program Studi'])
    
    biro = str(person['Biro/Bidang']).strip() if pd.notna(person['Biro/Bidang']) else ""
    jabatan_excel = proper_title_case(person['Jabatan']) if pd.notna(person['Jabatan']) else ""
    
    # Determine Jabatan text based on Biro/Bidang
    if biro.lower() == 'ketua umum':
        jabatan = 'Ketua Umum'
    elif biro in BIRO_BIDANG_MAP:
        jab_prefix = jabatan_excel if jabatan_excel else "Staf"
        jabatan = f"{jab_prefix} Bidang {BIRO_BIDANG_MAP[biro]}"
    else:
        jabatan = proper_title_case(f"{jabatan_excel} {biro}".strip())

    print(f"Mengekstrak data untuk: {nama}")
    
    try:
        doc = Document(template_doc_path)
    except Exception as e:
        print(f"Error reading Word template '{template_doc_path}': {e}")
        return False

    # Generate Nomor Surat
    now = datetime.datetime.now()
    roman_month = get_roman_month(now.month)
    year = now.year
    nomor_surat = f"No. {doc_number:03d}/SK/ARSC/{roman_month}/{year}"

    months_id = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
    tanggal_str = f"{now.day} {months_id[now.month]} {now.year}"
    
    for p in doc.paragraphs:
        text = p.text
        
        if "[TANGGAL_SURAT]" in text:
            set_paragraph_text(p, text.replace("[TANGGAL_SURAT]", tanggal_str))
            
        if "No." in text and "/SK/ARSC/" in text:
            set_paragraph_text(p, nomor_surat)
            
        if "[NAMA]" in text:
            set_paragraph_text(p, text.replace("[NAMA]", nama))
        if "[NIM]" in text:
            set_paragraph_text(p, text.replace("[NIM]", nim))
        if "[DEPARTEMEN]" in text:
            set_paragraph_text(p, text.replace("[DEPARTEMEN]", departemen))
        if "[PROGRAM_STUDI]" in text:
            set_paragraph_text(p, text.replace("[PROGRAM_STUDI]", program_studi))
        if "[JABATAN]" in text:
            set_paragraph_text(p, text.replace("[JABATAN]", jabatan))

    safe_nama = "".join([c for c in nama if c.isalpha() or c.isspace() or c == '-']).strip()
    output_filename = f"SK Aktif {safe_nama}.docx"
    output_path = os.path.join(BASE_DIR, "docx", output_filename)
    
    if os.path.exists(output_path):
        print(f" -> Lewati: {output_filename} sudah ada (Tidak menimpa file).")
        return False
        
    try:
        doc.save(output_path)
        print(f" -> Berhasil membuat: docx/{output_filename} (Nomor Surat: {doc_number:03d})")
        return True
    except Exception as e:
        print(f" -> Error saving Word file: {e}")
        return False

def get_current_surat_number(file_path=None, default_start=14):
    if file_path is None:
        file_path = os.path.join(BASE_DIR, "nomor_terakhir.txt")
    if os.path.exists(file_path):
        try:
            with open(file_path, 'r') as f:
                return int(f.read().strip())
        except:
            return default_start
    return default_start

def save_surat_number(num, file_path=None):
    if file_path is None:
        file_path = os.path.join(BASE_DIR, "nomor_terakhir.txt")
    try:
        with open(file_path, 'w') as f:
            f.write(str(num))
    except Exception as e:
        print(f"Peringatan: Gagal menyimpan nomor surat terakhir: {e}")

def main():
    parser = argparse.ArgumentParser(description="Generate Surat Keterangan Aktif dari Excel ke Word.")
    parser.add_argument("names", nargs="*", help="Nama lengkap anggota (bisa lebih dari satu).")
    parser.add_argument("--file", "-f", type=str, help="File teks berisi daftar nama (satu nama per baris).")
    parser.add_argument("--all", "-a", action="store_true", help="Generate SK untuk SEMUA anggota di Excel.")
    
    args = parser.parse_args()
    
    if not args.names and not args.file and not args.all:
        parser.print_help()
        print("\nContoh penggunaan:")
        print("  python generate_sk.py \"Nama Satu\" \"Nama Dua\"")
        print("  python generate_sk.py --file daftar_nama.txt")
        print("  python generate_sk.py --all")
        return

    excel_file = os.path.join(BASE_DIR, 'data', 'PENDATAAN ANGGOTA AKTIF ARSC PERIODE 2025_2026 .xlsx')
    try:
        df = pd.read_excel(excel_file)
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return

    names_to_process = []
    
    if args.all:
        names_to_process = df['Nama Lengkap'].dropna().tolist()
    else:
        if args.names:
            names_to_process.extend(args.names)
        if args.file:
            try:
                with open(args.file, 'r', encoding='utf-8') as f:
                    file_names = [line.strip() for line in f if line.strip()]
                    names_to_process.extend(file_names)
            except Exception as e:
                print(f"Error reading file {args.file}: {e}")
                return

    # Remove duplicates but keep order
    seen = set()
    names_to_process = [x for x in names_to_process if not (x.lower() in seen or seen.add(x.lower()))]

    current_surat_number = get_current_surat_number()
    success_count = 0
    
    excel_names = df['Nama Lengkap'].dropna().astype(str).tolist()
    excel_names_lower = {n.strip().lower(): n for n in excel_names}
    
    for target_name in names_to_process:
        target_lower = target_name.strip().lower()
        person_data = df[df['Nama Lengkap'].astype(str).str.strip().str.lower() == target_lower]
        
        if person_data.empty:
            matches = difflib.get_close_matches(target_lower, excel_names_lower.keys(), n=1, cutoff=0.6)
            if matches:
                best_match_lower = matches[0]
                actual_excel_name = excel_names_lower[best_match_lower]
                print(f"Data untuk nama '{target_name}' tidak ditemukan persis di Excel.")
                response = input(f"Apakah maksud Anda '{actual_excel_name}'? (y/n): ")
                if response.strip().lower() == 'y':
                    person_data = df[df['Nama Lengkap'].astype(str).str.strip().str.lower() == best_match_lower]
                else:
                    print(f"Melewati '{target_name}'.\n")
                    continue
            else:
                print(f"Data untuk nama '{target_name}' tidak ditemukan di Excel (typo terlalu jauh).\n")
                continue
        
        person = person_data.iloc[0]
        if create_docx_for_person(person, current_surat_number):
            current_surat_number += 1
            success_count += 1
            print() # empty line for readability
            
    # Save the updated number back to file
    save_surat_number(current_surat_number)
            
    print(f"Selesai! Berhasil membuat {success_count} dari {len(names_to_process)} permintaan surat keterangan.")

if __name__ == "__main__":
    main()
